import os
import tempfile
from typing import List, Optional

from PyPDF2 import PdfReader
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.prompts import PromptTemplate
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.messages import AIMessage, HumanMessage

from utils import clear_directory

CHROMA_PERSIST_DIR = "./chroma_db"

def get_text_from_pdf(pdf_path: str) -> str:
    text = ""
    with open(pdf_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += f"\\n--- Page {i + 1} ---\\n{page_text}"
    return text

def get_text_from_docx(docx_path: str) -> str:
    doc = docx.Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        if para.text:
            text += para.text + "\\n"
    return text

def process_document(file_name: str, file_path: str) -> List[str]:
    ext = os.path.splitext(file_name)[1].lower()
    if ext == ".pdf":
        return get_text_from_pdf(file_path)
    elif ext == ".docx":
        return get_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def chunk_text(text: str, source: str) -> List[dict]:
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    
    # Store source metadata
    documents = []
    for chunk in chunks:
        documents.append({
            "page_content": chunk,
            "metadata": {"source": source}
        })
    return documents

def get_vector_store():
    # Only using Gemini embeddings
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = Chroma(
        embedding_function=embeddings,
        persist_directory=CHROMA_PERSIST_DIR
    )
    return vector_store

def add_documents_to_db(documents: list, vector_store):
    from langchain_core.documents import Document
    docs = [Document(page_content=d["page_content"], metadata=d["metadata"]) for d in documents]
    vector_store.add_documents(docs)

def clear_vector_db():
    vector_store = get_vector_store()
    vector_store.delete_collection()
    clear_directory(CHROMA_PERSIST_DIR)

def get_chat_chain():
    llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro", temperature=0.3)
    
    prompt_template = """
    You are an intelligent AI assistant. Think of yourself as an incredibly smart student 
    having an open-book test using the teacher's materials.
    
    Answer the user's question based ONLY on the provided context.
    If the answer is not contained within the provided context, gracefully state that 
    you do not have enough information to answer based on the given documents.
    
    Cite your sources implicitly based on the metadata (e.g., "According to the uploaded syllabus...").
    
    <context>
    {context}
    </context>
    
    Chat History: {chat_history}
    
    Question: {input}
    
    Assistant:"""
    
    prompt = PromptTemplate(
        template=prompt_template,
        input_variables=["context", "chat_history", "input"]
    )
    
    document_chain = create_stuff_documents_chain(llm, prompt)
    
    vector_store = get_vector_store()
    retriever = vector_store.as_retriever(search_kwargs={"k": 5})
    
    retrieval_chain = create_retrieval_chain(retriever, document_chain)
    return retrieval_chain

def format_chat_history(messages):
    formatted = []
    for msg in messages:
        if msg["role"] == "user":
            formatted.append(f"Human: {msg['content']}")
        else:
            formatted.append(f"AI: {msg['content']}")
    return "\\n".join(formatted)
