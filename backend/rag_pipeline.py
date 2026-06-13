import os
import tempfile
import json
from typing import List, Optional

from PyPDF2 import PdfReader
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.prompts import PromptTemplate
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.messages import AIMessage, HumanMessage
import anthropic

from utils import clear_directory

CHROMA_PERSIST_DIR = "./chroma_db"

def get_text_from_pdf(pdf_path: str) -> str:
    text = ""
    with open(pdf_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += f"\n--- Page {i + 1} ---\n{page_text}"
    return text

def get_text_from_docx(docx_path: str) -> str:
    doc = docx.Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        if para.text:
            text += para.text + "\n"
    return text

def process_document(file_name: str, file_path: str) -> List[str]:
    ext = os.path.splitext(file_name)[1].lower()
    if ext == ".pdf":
        return get_text_from_pdf(file_path)
    elif ext == ".docx":
        return get_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def chunk_text(text: str, source: str) -> List[dict]:
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    
    # Store source metadata
    documents = []
    for chunk in chunks:
        documents.append({
            "page_content": chunk,
            "metadata": {"source": source}
        })
    return documents

def get_vector_store():
    # Only using Gemini embeddings
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = Chroma(
        embedding_function=embeddings,
        persist_directory=CHROMA_PERSIST_DIR
    )
    return vector_store

def add_documents_to_db(documents: list, vector_store):
    from langchain_core.documents import Document
    docs = [Document(page_content=d["page_content"], metadata=d["metadata"]) for d in documents]
    vector_store.add_documents(docs)

def clear_vector_db():
    vector_store = get_vector_store()
    vector_store.delete_collection()
    clear_directory(CHROMA_PERSIST_DIR)

def get_chat_chain(intent: str = "question"):
    # Using gemini-1.5-pro for better reasoning in the main chat
    llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro", temperature=0.3, streaming=True)
    
    # Customize prompt based on intent for a "smoother" feel
    if intent == "greeting":
        system_prompt = "You are a friendly and professional engineering assistant from Partner AI. Greet the user warmly and ask what they would like to build or learn today."
    elif intent == "task":
        system_prompt = "You are an expert project manager. Help the user refine their task and offer to break it down using the GSD tool."
    else:
        system_prompt = "You are an intelligent engineering assistant. Answer based on the provided context. Be concise and technical."

    prompt_template = f"""
    {{system_prompt}}
    
    Context:
    {{context}}
    
    Chat History:
    {{chat_history}}
    
    Human: {{input}}
    Assistant:"""
    
    prompt = PromptTemplate(
        template=prompt_template,
        input_variables=["context", "chat_history", "input"],
        partial_variables={"system_prompt": system_prompt}
    )
    
    document_chain = create_stuff_documents_chain(llm, prompt)
    vector_store = get_vector_store()
    retriever = vector_store.as_retriever(search_kwargs={"k": 5})
    
    return create_retrieval_chain(retriever, document_chain)

def format_chat_history(messages):
    formatted = []
    for msg in messages:
        if msg["role"] == "user":
            formatted.append(f"Human: {msg['content']}")
        else:
            formatted.append(f"AI: {msg['content']}")
    return "\n".join(formatted)

async def generate_live_outline(text_chunk: str) -> list:
    llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)
    prompt = f"Extract a maximum of 3 highly important, extremely concise bullet points (1 sentence each) from the following live lecture chunk:\n\n{text_chunk}\n\nFormat your output strictly as a JSON array of strings, e.g. [\"point 1\", \"point 2\"]. Return NOTHING else."
    
    try:
        response = await llm.ainvoke(prompt)
        text = response.content.replace('```json', '').replace('```', '').strip()
        return json.loads(text)
    except Exception as e:
        print(f"Error generating outline: {e}")
        return [f"Key point recorded."]

async def generate_live_quiz(outline_history: list) -> list:
    topics = "\n".join(outline_history[-5:]) # Last 5 outline points
    
    system_prompt = """You are an expert educational quiz designer. 
Given recent lecture topics, generate a clear, accurate, and engaging quiz to test the student's understanding.
Respond ONLY with a valid JSON array — no markdown, no explanation, no backticks.
JSON structure:
[
  {
    "question": "Question text?",
    "options": ["A) ...", "B) ...", "C) ...", "D) ..."],
    "answer": "A) ..."
  }
]"""

    user_prompt = f"""Generate a 2-question multiple choice quiz from these lecture topics.
Make questions test real understanding, not trivia. Use simple language.

TOPICS:
{topics}"""

    try:
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            print("Anthropic API key missing")
            return []
            
        client = anthropic.AsyncAnthropic(api_key=api_key)
        response = await client.messages.create(
            model="claude-3-5-sonnet-20240620",
            max_tokens=1000,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )
        raw = response.content[0].text.strip()
        
        # Strip any accidental markdown fences
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
                
        return json.loads(raw.strip())
    except Exception as e:
        print(f"Error generating quiz with Anthropic: {e}")
        return []

async def generate_live_recap(outline_history: list) -> str:
    llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)
    topics = "\n".join(outline_history)
    prompt = f"Write a comprehensive, engaging 2-3 sentence summary of today's class based on these lecture outlines:\n\n{topics}"
    
    try:
        response = await llm.ainvoke(prompt)
        return response.content.strip()
    except Exception:
        return "Thank you for attending today's session."

async def categorize_intent(message: str) -> str:
    llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0)
    prompt = f"Categorize the following user message into one of these categories: 'question', 'greeting', 'feedback', 'task', or 'unknown'. Message: {message}. Return only the category name."
    try:
        response = await llm.ainvoke(prompt)
        return response.content.strip().lower()
    except Exception:
        return "unknown"
